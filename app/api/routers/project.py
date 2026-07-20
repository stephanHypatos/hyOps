
import os
import re
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from uuid import UUID
from datetime import datetime
from sqlmodel import select
from enum import Enum
from app.database.models import (
    Project, ProjectStakeholder, ProjectUsecase, ProjectFeature,
    User, Usecase, Feature, Capability, DocumentTemplate,
    GeneratedDocument, DocumentStatus, Organization
)
from app.database.session import SessionDep
from app.api.schemas.project import (
    ProjectCreate, ProjectRead, ProjectUpdate,
    StakeholderRead, StakeholderReadEnriched, StakeholderAssign,
    UsecaseAssign, UsecaseRead, ProjectFeatureAssign, GenerateFromTemplatesRequest
)
from app.api.schemas.feature import FeatureReadWithCapability

from app.config import basedir


router = APIRouter(prefix="/project", tags=["Project"])


# ===================== Project CRUD Endpoints =====================

@router.get("/projects", response_model=list[ProjectRead])
async def get_all_projects(session: SessionDep):
    result = await session.execute(select(Project))
    return result.scalars().all()


@router.get("/", response_model=ProjectRead)
async def get_project(id: UUID, session: SessionDep):
    project = await session.get(Project, id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project


@router.post("/", response_model=dict)
async def create_project(data: ProjectCreate, session: SessionDep):
    new_project = Project(**data.model_dump())
    session.add(new_project)
    await session.commit()
    await session.refresh(new_project)
    return {"id": new_project.id}


@router.patch("/", response_model=ProjectRead)
async def update_project(id: UUID, data: ProjectUpdate, session: SessionDep):
    update_data = data.model_dump(exclude_none=True)
    if not update_data:
        raise HTTPException(status_code=400, detail="No data to update")
    
    project = await session.get(Project, id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    update_data["updated_at"] = datetime.now()
    project.sqlmodel_update(update_data)
    
    session.add(project)
    await session.commit()
    await session.refresh(project)
    return project


@router.delete("/")
async def delete_project(id: UUID, session: SessionDep):
    project = await session.get(Project, id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    await session.delete(project)
    await session.commit()
    return {"detail": "Deleted"}


# ===================== Stakeholder M2M Endpoints =====================

@router.get("/{project_id}/stakeholders", response_model=list[StakeholderReadEnriched])
async def get_project_stakeholders(project_id: UUID, session: SessionDep):
    """
    Get stakeholders with enriched user details (name, email).
    Single query JOIN eliminates N+1 lookups on the frontend.
    """
    result = await session.execute(
        select(ProjectStakeholder, User)
        .join(User, ProjectStakeholder.user_id == User.id)
        .where(ProjectStakeholder.project_id == project_id)
    )
    rows = result.all()
    
    return [
        StakeholderReadEnriched(
            project_id=ps.project_id,
            user_id=ps.user_id,
            role=ps.role,
            first_name=user.first_name,
            last_name=user.last_name,
            email=user.email
        )
        for ps, user in rows
    ]


@router.post("/{project_id}/stakeholder", response_model=StakeholderRead)
async def assign_stakeholder(project_id: UUID, data: StakeholderAssign, session: SessionDep):
    # Validate user exists
    user = await session.get(User, data.user_id)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Validate project exists
    project = await session.get(Project, project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Check for duplicate
    existing = await session.get(ProjectStakeholder, (project_id, data.user_id))
    if existing:
        raise HTTPException(status_code=400, detail="User is already a stakeholder in this project")
    
    link = ProjectStakeholder(project_id=project_id, **data.model_dump())
    session.add(link)
    await session.commit()
    await session.refresh(link)
    return link


@router.delete("/{project_id}/stakeholder/{user_id}")
async def remove_stakeholder(project_id: UUID, user_id: UUID, session: SessionDep):
    link = await session.get(ProjectStakeholder, (project_id, user_id))
    if not link:
        raise HTTPException(status_code=404, detail="Stakeholder not found in this project")
    await session.delete(link)
    await session.commit()
    return {"detail": "Stakeholder removed"}


# ===================== Usecase M2M Endpoints =====================

@router.get("/{project_id}/usecases", response_model=list[UsecaseRead])
async def get_project_usecases(project_id: UUID, session: SessionDep):
    result = await session.execute(
        select(ProjectUsecase).where(ProjectUsecase.project_id == project_id)
    )
    return result.scalars().all()


@router.post("/{project_id}/usecase")
async def assign_usecase(project_id: UUID, data: UsecaseAssign, session: SessionDep):
    # Validate usecase exists
    usecase = await session.get(Usecase, data.usecase_id)
    if not usecase:
        raise HTTPException(status_code=404, detail="Usecase not found")
    
    # Check for duplicate
    existing = await session.get(ProjectUsecase, (project_id, data.usecase_id))
    if existing:
        raise HTTPException(status_code=400, detail="Usecase is already in this project")
    
    link = ProjectUsecase(project_id=project_id, **data.model_dump())
    session.add(link)
    await session.commit()
    await session.refresh(link)
    return link


@router.delete("/{project_id}/usecase/{usecase_id}")
async def remove_usecase(project_id: UUID, usecase_id: UUID, session: SessionDep):
    link = await session.get(ProjectUsecase, (project_id, usecase_id))
    if not link:
        raise HTTPException(status_code=404, detail="Usecase not found in this project")
    await session.delete(link)
    await session.commit()
    return {"detail": "Usecase removed from project"}


# ===================== Custom Feature M2M Endpoints =====================

@router.get("/{project_id}/features", response_model=list[FeatureReadWithCapability])
async def get_project_features(project_id: UUID, session: SessionDep):
    """
    Get custom features attached directly to this project (independent of use
    cases), enriched with capability details. Single JOIN avoids N+1 lookups.
    """
    result = await session.execute(
        select(Feature, Capability)
        .join(Capability, Feature.capability_id == Capability.id)
        .join(ProjectFeature, ProjectFeature.feature_id == Feature.id)
        .where(ProjectFeature.project_id == project_id)
    )
    rows = result.all()

    return [
        FeatureReadWithCapability(
            id=feature.id,
            capability_id=feature.capability_id,
            name=feature.name,
            service_description=feature.service_description,
            deliverables=feature.deliverables,
            scope_type=feature.scope_type,
            owner_id=feature.owner_id,
            scoping_questionnaire=feature.scoping_questionnaire,
            reference_documentation=feature.reference_documentation,
            included_in_ootb=feature.included_in_ootb,
            default_enabled=feature.default_enabled,
            active=feature.active,
            multiple_value=feature.multiple_value,
            requirements=feature.requirements or [],
            created_at=feature.created_at,
            updated_at=feature.updated_at,
            capability_name=capability.name,
            capability_contract=capability.contract,
        )
        for feature, capability in rows
    ]


@router.post("/{project_id}/feature")
async def assign_feature(project_id: UUID, data: ProjectFeatureAssign, session: SessionDep):
    # Validate project exists
    project = await session.get(Project, project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Validate feature exists
    feature = await session.get(Feature, data.feature_id)
    if not feature:
        raise HTTPException(status_code=404, detail="Feature not found")

    # Check for duplicate
    existing = await session.get(ProjectFeature, (project_id, data.feature_id))
    if existing:
        raise HTTPException(status_code=400, detail="Feature is already attached to this project")

    link = ProjectFeature(project_id=project_id, feature_id=data.feature_id)
    session.add(link)
    await session.commit()
    await session.refresh(link)
    return link


@router.delete("/{project_id}/feature/{feature_id}")
async def remove_feature(project_id: UUID, feature_id: UUID, session: SessionDep):
    link = await session.get(ProjectFeature, (project_id, feature_id))
    if not link:
        raise HTTPException(status_code=404, detail="Feature not attached to this project")
    await session.delete(link)
    await session.commit()
    return {"detail": "Feature removed from project"}





# ===================== Document Generation Endpoint =====================




# ===================== Helper Functions =====================

FEATURE_SKIP = {
    "id", "capability_id", "owner_id", "created_at", "updated_at",
    "multiple_value", "scoping_questionnaire",
    "included_in_ootb", "default_enabled", "active", "capability", "owner",
    "scope_specifications", "cost_drivers", "feature_efforts", "usecases"
}


PROJECT_SKIP = {
    "id", "created_at", "updated_at",
    "customer", "deal_winner", "stakeholders",
    "linked_usecases", "primary_usecase",
    "generated_documents", "erp_connectors", "hystudio_companies"
}
ORG_SKIP = {
    "id", "created_at", "updated_at",
    "users", "projects", "teams_groups", "slack_channels",
    "metabase_groups", "hystudio_companies", "erp_systems"
}
USER_SKIP = {
    "id", "created_at", "updated_at",
    "organization", "subtype", "languages", "skills",
    "won_projects", "owned_features", "created_templates",
    "projects", "teams_groups", "slack_channels",
    "metabase_groups", "hystudio_companies"
}
USECASE_SKIP = {"id", "features", "projects"}



def _serialize_value(val):
    if val is None:
        return ""
    # FIX: Check for Enum first and return its actual value (e.g., "pilot" instead of "ProjectType.pilot")
    if isinstance(val, Enum):
        return val.value
    if hasattr(val, 'isoformat'):
        return val.isoformat()
    if isinstance(val, bool):
        return str(val).lower()
    if isinstance(val, (str, int, float)):
        return val
    if isinstance(val, list):
        return [_serialize_value(v) for v in val]
    if isinstance(val, dict):
        return {k: _serialize_value(v) for k, v in val.items()}
    return str(val)


def _empty_fields(model_cls, skip_set):
    """
    Same keys as _extract_fields would produce, but all blank.

    Used for optional relations (e.g. a project with no partner) so that
    {{partner.key}} renders as an empty string instead of leaking the literal
    placeholder text into a customer-facing document.
    """
    fields = getattr(model_cls, 'model_fields', None) or getattr(model_cls, '__sqlmodel_fields__', {})
    return {col: "" for col in fields if col not in skip_set}


def _extract_fields(obj, skip_set):
    result = {}
    # Works with both Pydantic v1 and v2
    fields = getattr(type(obj), 'model_fields', None) or getattr(type(obj), '__sqlmodel_fields__', {})
    for col in fields:
        if col in skip_set:
            continue
        result[col] = _serialize_value(getattr(obj, col, None))
    return result



def build_render_context(project: Project) -> dict:
    """Build nested dict from Project ORM for {{project.name}} style placeholders."""
    context = {
        "project": _extract_fields(project, PROJECT_SKIP),
        "customer": (
            _extract_fields(project.customer, ORG_SKIP) if project.customer
            else _empty_fields(Organization, ORG_SKIP)
        ),
        "partner": (
            _extract_fields(project.partner, ORG_SKIP) if project.partner
            else _empty_fields(Organization, ORG_SKIP)
        ),
        "deal_winner": (
            _extract_fields(project.deal_winner, USER_SKIP) if project.deal_winner
            else _empty_fields(User, USER_SKIP)
        ),
        "primary_usecase": (
            _extract_fields(project.primary_usecase, USECASE_SKIP) if project.primary_usecase
            else _empty_fields(Usecase, USECASE_SKIP)
        ),
        "stakeholders": [],
        "features": [], # NEW: Add features list for iteration
    }

    if project.stakeholders:
        for user in project.stakeholders:
            context["stakeholders"].append({
                "first_name": user.first_name,
                "last_name": user.last_name,
                "email": getattr(user, 'email', ''),
                "phone": getattr(user, 'phone', ''),
                # _serialize_value unwraps the enum — str() would leak "UserType.internal"
                "type": _serialize_value(getattr(user, 'type', '')),
                "role": _serialize_value(getattr(user, 'role', '')),
                "subtype": str(getattr(user.subtype, 'name', '')) if hasattr(user, 'subtype') and user.subtype else '',
            })

    # Extract features from the primary usecase, then merge in any custom
    # features attached directly to the project. De-duplicate by feature id so
    # a feature that is both in the usecase and attached directly appears once.
    seen_feature_ids = set()

    if project.primary_usecase and project.primary_usecase.features:
        for feature in project.primary_usecase.features:
            seen_feature_ids.add(feature.id)
            context["features"].append(_extract_fields(feature, FEATURE_SKIP))

    if project.custom_features:
        for feature in project.custom_features:
            if feature.id in seen_feature_ids:
                continue
            seen_feature_ids.add(feature.id)
            context["features"].append(_extract_fields(feature, FEATURE_SKIP))

    return context




LOOP_PATTERN = r'\{\{#(\w+)\}\}(.*?)\{\{/\1\}\}'


def _render_loops(template: str, context: dict) -> str:
    """
    Render {{#list}} ... {{/list}} blocks.

    Recurses, so loops can be nested — e.g. a {{#requirements}} block inside a
    {{#features}} block. The inner loop is resolved against the current outer
    item, which is how per-feature requirements get rendered:

        {{#features}}
        ### {{feature.name}}
        {{#requirements}}
        - {{requirement.requirement}}: {{requirement.solution}}
        {{/requirements}}
        {{/features}}
    """
    def loop_replacer(match):
        list_key = match.group(1)
        inner_template = match.group(2).strip()
        items = _resolve_dot_path(list_key, context)

        if not isinstance(items, list):
            return ""

        # Handle singular key (e.g. features -> feature)
        singular_key = list_key[:-1] if list_key.endswith('s') else list_key
        rendered_items = []

        for item in items:
            if not isinstance(item, dict):
                continue

            # Resolve any nested loops first, scoped to this item
            item_rendered = _render_loops(inner_template, item)

            # Then replace scalars like {{feature.name}} or bare {{name}}
            for k, v in item.items():
                if isinstance(v, dict):
                    continue
                if isinstance(v, list):
                    # A list of objects is rendered by a nested loop, not inline
                    if v and isinstance(v[0], dict):
                        continue
                    val = ", ".join(str(x) for x in v)
                else:
                    val = str(v) if v is not None else ""
                item_rendered = item_rendered.replace(f'{{{{{singular_key}.{k}}}}}', val)
                item_rendered = item_rendered.replace(f'{{{{{k}}}}}', val)

            rendered_items.append(item_rendered)

        return "\n".join(rendered_items)

    return re.sub(LOOP_PATTERN, loop_replacer, template, flags=re.DOTALL)


def replace_placeholders(markdown_content: str, context: dict) -> str:
    """Pure regex replacement for {{project.name}}, {{customer.email}}, and #loops."""

    # 1. Handle list loops first (supports nesting)
    markdown_content = _render_loops(markdown_content, context)

    # 2. Handle simple variables (existing logic + singular list fallback)
    def replacer(match):
        var_path = match.group(1).strip()
        value = _resolve_dot_path(var_path, context)
        
        if value is None:
            # Fallback: bare key like {{name}} → try under "project"
            if "." not in var_path:
                value = _resolve_dot_path(f"project.{var_path}", context)
                
        # NEW: Handle singular.key outside of loops (e.g., {{feature.name}} -> joins all feature names)
        if value is None and "." in var_path:
            parts = var_path.split(".", 1)
            singular = parts[0]      # e.g., "feature"
            nested_key = parts[1]    # e.g., "name"
            
            # Try to find the plural list in context (e.g., "feature" -> "features")
            plural = singular + "s" if not singular.endswith("s") else singular
            list_items = _resolve_dot_path(plural, context)
            
            if isinstance(list_items, list):
                extracted = []
                for item in list_items:
                    if isinstance(item, dict) and nested_key in item:
                        if item[nested_key] is not None:
                            extracted.append(str(item[nested_key]))
                if extracted:
                    value = ", ".join(extracted)
                    
        if value is None:
            return match.group(0)
            
        if isinstance(value, list):
            if len(value) > 0 and isinstance(value[0], dict):
                lines = []
                for item in value:
                    lines.append("- " + ", ".join(f"{k}: {v}" for k, v in item.items()))
                return "\n".join(lines)
            return ", ".join(str(v) for v in value)
        if isinstance(value, dict):
            return ", ".join(f"{k}: {v}" for k, v in value.items())
        return str(value)

    pattern = r'\{\{([\w.]+)\}\}'
    return re.sub(pattern, replacer, markdown_content)


def _resolve_dot_path(path: str, context: dict):
    current = context
    for part in path.split("."):
        if isinstance(current, dict) and part in current:
            current = current[part]
        else:
            return None
    return current


def save_generated_file(output_dir: str, filename: str, content: str, file_format: str) -> str:
    """Save rendered content to disk as .md or .docx"""
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, filename)

    if file_format == "docx":
        from docx import Document
        doc = Document()
        for line in content.split('\n'):
            doc.add_paragraph(line)
        doc.save(file_path)
    else:
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

    return file_path




# ===================== Template-Based Generation Endpoints =====================

def _media_type(file_format: str) -> str:
    if file_format == "docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if file_format == "md":
        return "text/markdown"
    return "application/octet-stream"


def _version_dir(project_id: UUID, template_id: UUID) -> str:
    """One directory per template, holding every version of that document."""
    return os.path.join(basedir, 'generated_documents', str(project_id), str(template_id))


def _serialize_version(doc: GeneratedDocument, template: DocumentTemplate | None) -> dict:
    return {
        "id": str(doc.id),
        "template_id": str(doc.template_id),
        "template_name": template.name if template else "Unknown",
        "template_type": (
            template.type.value if template and hasattr(template.type, 'value')
            else (str(template.type) if template else "")
        ),
        "template_version": doc.template_version,
        "version_no": doc.version_no,
        "file_format": doc.file_format,
        "status": doc.status.value if hasattr(doc.status, 'value') else str(doc.status),
        "created_at": doc.created_at.isoformat(),
    }


@router.get("/{project_id}/generated-documents")
async def get_generated_documents(project_id: UUID, session: SessionDep):
    """
    List the latest version of each generated document for this project.
    Backed by the generated_document version ledger.
    """
    result = await session.execute(
        select(GeneratedDocument)
        .where(GeneratedDocument.project_id == project_id)
        .order_by(GeneratedDocument.version_no.desc())
    )
    docs = result.scalars().all()

    # Rows are ordered newest-first, so the first hit per template is the latest
    latest_by_template: dict[UUID, GeneratedDocument] = {}
    for doc in docs:
        latest_by_template.setdefault(doc.template_id, doc)

    results = []
    for template_id, doc in latest_by_template.items():
        template = await session.get(DocumentTemplate, template_id)
        payload = _serialize_version(doc, template)
        payload["total_versions"] = sum(1 for d in docs if d.template_id == template_id)
        results.append(payload)

    results.sort(key=lambda r: r["created_at"], reverse=True)
    return results


@router.get("/{project_id}/document-versions/{template_id}")
async def get_document_versions(project_id: UUID, template_id: UUID, session: SessionDep):
    """Full version history for one document, newest first."""
    template = await session.get(DocumentTemplate, template_id)

    result = await session.execute(
        select(GeneratedDocument)
        .where(GeneratedDocument.project_id == project_id)
        .where(GeneratedDocument.template_id == template_id)
        .order_by(GeneratedDocument.version_no.desc())
    )
    return [_serialize_version(doc, template) for doc in result.scalars().all()]


@router.post("/{project_id}/generate-from-templates")
async def generate_from_templates(
    project_id: UUID,
    data: GenerateFromTemplatesRequest,
    session: SessionDep
):
    """
    Generate documents from selected templates.

    Never overwrites: each run appends a new version (v1, v2, v3 …) both on disk
    and in the generated_document ledger, so earlier versions stay downloadable.
    """
    project = await session.get(Project, project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    context = build_render_context(project)
    results = []

    for template_id in data.template_ids:
        template = await session.get(DocumentTemplate, template_id)

        # VALIDATION: Skip if template doesn't exist OR is not active
        if not template or not template.is_active:
            continue

        # 1. Replace placeholders
        rendered = replace_placeholders(template.markdown_content, context)

        # 2. Next version number for this project + template
        latest = (await session.execute(
            select(GeneratedDocument)
            .where(GeneratedDocument.project_id == project_id)
            .where(GeneratedDocument.template_id == template_id)
            .order_by(GeneratedDocument.version_no.desc())
        )).scalars().first()
        version_no = (latest.version_no + 1) if latest else 1

        # 3. Save to disk as {project_id}/{template_id}/v{n}.{format}
        file_format = template.file_format or 'md'
        file_path = save_generated_file(
            _version_dir(project_id, template_id),
            f"v{version_no}.{file_format}",
            rendered,
            file_format,
        )

        # 4. Record the version
        doc = GeneratedDocument(
            project_id=project_id,
            template_id=template_id,
            version_no=version_no,
            document_type=template.type.value if hasattr(template.type, 'value') else str(template.type),
            markdown_generated=rendered,
            file_path=file_path,
            file_format=file_format,
            template_version=template.version,
            status=DocumentStatus.final,
        )
        session.add(doc)
        await session.commit()
        await session.refresh(doc)

        payload = _serialize_version(doc, template)
        payload["total_versions"] = version_no
        results.append(payload)

    # Optional: Handle case where all requested templates were inactive/not found
    if not results and data.template_ids:
        raise HTTPException(status_code=400, detail="All selected templates are either inactive or do not exist.")

    return {
        "detail": f"Generated {len(results)} document(s) successfully",
        "documents": results
    }


@router.get("/{project_id}/download-template/{template_id}")
async def download_template_document(
    project_id: UUID,
    template_id: UUID,
    session: SessionDep
):
    """Download the latest version of a generated document."""
    doc = (await session.execute(
        select(GeneratedDocument)
        .where(GeneratedDocument.project_id == project_id)
        .where(GeneratedDocument.template_id == template_id)
        .order_by(GeneratedDocument.version_no.desc())
    )).scalars().first()

    if not doc:
        raise HTTPException(status_code=404, detail="No generated document found. Generate it first.")

    return _version_file_response(doc, await session.get(DocumentTemplate, template_id))


@router.get("/{project_id}/download-version/{version_id}")
async def download_document_version(
    project_id: UUID,
    version_id: UUID,
    session: SessionDep
):
    """Download one specific version of a generated document."""
    doc = await session.get(GeneratedDocument, version_id)
    if not doc or doc.project_id != project_id:
        raise HTTPException(status_code=404, detail="Version not found for this project")

    return _version_file_response(doc, await session.get(DocumentTemplate, doc.template_id))


def _version_file_response(doc: GeneratedDocument, template: DocumentTemplate | None) -> FileResponse:
    if not os.path.exists(doc.file_path):
        raise HTTPException(
            status_code=404,
            detail=f"File for v{doc.version_no} is missing on disk.",
        )

    # Create a human-readable download filename, e.g. Statement_of_Work_v3.md
    base = template.name if template else "document"
    safe_name = re.sub(r'[^\w]', '_', base)[:40]
    download_name = f"{safe_name}_v{doc.version_no}.{doc.file_format}"

    return FileResponse(
        path=doc.file_path,
        media_type=_media_type(doc.file_format),
        filename=download_name,
    )